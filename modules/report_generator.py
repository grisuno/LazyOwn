#!/usr/bin/env python3
"""
modules/report_generator.py
============================
Auto-generates a structured penetration test report from LazyOwn session data.

Sources read (all optional — missing files are skipped gracefully):
  sessions/policy_facts.json    FactStore structured data (hosts, services, vulns)
  sessions/events.jsonl         Event timeline
  sessions/objectives.jsonl     Attack objectives with status
  sessions/credentials.txt      Captured credentials
  sessions/plan.txt             VulnBot/LLM attack plan
  sessions/sessionLazyOwn.json  Session state snapshot

Output:
  sessions/report_<YYYYMMDD_HHMMSS>.md   (default)
  sessions/report_<YYYYMMDD_HHMMSS>.docx (when output ends in .docx or formats includes "docx")

New capabilities:
  CVSSv3Calculator  -- compute CVSS v3.1 base scores from vector strings
  DOCXExporter      -- export Markdown reports to .docx (requires python-docx)

Usage:
    from modules.report_generator import ReportGenerator

    rg   = ReportGenerator()
    path = rg.generate()
    print(f"Report: {path}")

    # Generate both md and docx simultaneously:
    paths = rg.generate(formats=["md", "docx"])

    # CVSS score:
    from modules.report_generator import CVSSv3Calculator
    score, severity = CVSSv3Calculator().calculate("AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H")

    # CLI:
    python3 modules/report_generator.py [--sessions PATH] [--output PATH] [--quiet]
"""
from __future__ import annotations

import json
import logging
import math
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

log = logging.getLogger("report_generator")


# ---------------------------------------------------------------------------
# CVSSv3Calculator (Single Responsibility: CVSS v3.1 scoring only)
# ---------------------------------------------------------------------------


class CVSSv3Calculator:
    """
    Compute CVSS v3.1 base scores without any external dependencies.

    Implements the formula from the CVSS v3.1 specification
    (https://www.first.org/cvss/v3.1/specification-document).

    Usage::

        calc = CVSSv3Calculator()
        score, severity = calc.calculate("AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H")
        # (9.8, "Critical")

        score, severity = calc.from_nvd_metrics(
            attackVector="NETWORK",
            attackComplexity="LOW",
            privilegesRequired="NONE",
            userInteraction="NONE",
            scope="UNCHANGED",
            confidentialityImpact="HIGH",
            integrityImpact="HIGH",
            availabilityImpact="HIGH",
        )
    """

    # ------------------------------------------------------------------
    # CVSS v3.1 metric weight tables
    # ------------------------------------------------------------------

    _AV  = {"N": 0.85, "A": 0.62, "L": 0.55, "P": 0.20}
    _AC  = {"L": 0.77, "H": 0.44}
    _PR_UNCHANGED  = {"N": 0.85, "L": 0.62, "H": 0.27}
    _PR_CHANGED    = {"N": 0.85, "L": 0.68, "H": 0.50}
    _UI  = {"N": 0.85, "R": 0.62}
    _CIA = {"N": 0.00, "L": 0.22, "H": 0.56}

    # Abbreviated metric value -> canonical abbreviation mapping
    _AV_MAP  = {
        "N": "N", "NETWORK": "N",
        "A": "A", "ADJACENT": "A", "ADJACENT_NETWORK": "A",
        "L": "L", "LOCAL": "L",
        "P": "P", "PHYSICAL": "P",
    }
    _AC_MAP  = {"L": "L", "LOW": "L", "H": "H", "HIGH": "H"}
    _PR_MAP  = {"N": "N", "NONE": "N", "L": "L", "LOW": "L", "H": "H", "HIGH": "H"}
    _UI_MAP  = {"N": "N", "NONE": "N", "R": "R", "REQUIRED": "R"}
    _S_MAP   = {"U": "U", "UNCHANGED": "U", "C": "C", "CHANGED": "C"}
    _CIA_MAP = {"N": "N", "NONE": "N", "L": "L", "LOW": "L", "H": "H", "HIGH": "H"}

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def calculate(self, vector_string: str) -> Tuple[float, str]:
        """
        Compute the CVSS v3.1 base score from a vector string.

        Accepts both full (CVSS:3.1/AV:N/...) and short (AV:N/...) forms.

        Returns
        -------
        Tuple[float, str]
            (base_score rounded to 1 decimal, severity label)
        """
        metrics = self._parse_vector(vector_string)
        return self._compute(
            av=metrics["AV"],
            ac=metrics["AC"],
            pr=metrics["PR"],
            ui=metrics["UI"],
            s=metrics["S"],
            c=metrics["C"],
            i=metrics["I"],
            a=metrics["A"],
        )

    def from_nvd_metrics(
        self,
        attackVector: str,
        attackComplexity: str,
        privilegesRequired: str,
        userInteraction: str,
        scope: str,
        confidentialityImpact: str,
        integrityImpact: str,
        availabilityImpact: str,
    ) -> Tuple[float, str]:
        """
        Compute the CVSS v3.1 base score from individual NVD metric names.

        All parameters accept both full names ("NETWORK") and abbreviations
        ("N").

        Returns
        -------
        Tuple[float, str]
            (base_score rounded to 1 decimal, severity label)
        """
        av = self._AV_MAP.get(attackVector.upper(), "N")
        ac = self._AC_MAP.get(attackComplexity.upper(), "L")
        pr = self._PR_MAP.get(privilegesRequired.upper(), "N")
        ui = self._UI_MAP.get(userInteraction.upper(), "N")
        s  = self._S_MAP.get(scope.upper(), "U")
        c  = self._CIA_MAP.get(confidentialityImpact.upper(), "N")
        i  = self._CIA_MAP.get(integrityImpact.upper(), "N")
        a  = self._CIA_MAP.get(availabilityImpact.upper(), "N")
        return self._compute(av=av, ac=ac, pr=pr, ui=ui, s=s, c=c, i=i, a=a)

    # ------------------------------------------------------------------
    # Internal computation
    # ------------------------------------------------------------------

    def _compute(
        self,
        av: str, ac: str, pr: str, ui: str,
        s: str, c: str, i: str, a: str,
    ) -> Tuple[float, str]:
        """Apply the CVSS v3.1 base score formula."""
        scope_changed = s == "C"

        av_val  = self._AV[av]
        ac_val  = self._AC[ac]
        pr_val  = (self._PR_CHANGED if scope_changed else self._PR_UNCHANGED)[pr]
        ui_val  = self._UI[ui]
        c_val   = self._CIA[c]
        i_val   = self._CIA[i]
        a_val   = self._CIA[a]

        # Exploitability sub-score
        exploitability = 8.22 * av_val * ac_val * pr_val * ui_val

        # Impact sub-score
        iss = 1.0 - (1.0 - c_val) * (1.0 - i_val) * (1.0 - a_val)

        if iss <= 0.0:
            return (0.0, "None")

        if scope_changed:
            impact = 7.52 * (iss - 0.029) - 3.25 * ((iss - 0.02) ** 15)
        else:
            impact = 6.42 * iss

        raw = impact + exploitability

        if scope_changed:
            raw = min(raw, 10.0)

        base_score = self._roundup(raw)
        severity   = self._severity(base_score)
        return (base_score, severity)

    @staticmethod
    def _roundup(value: float) -> float:
        """CVSS v3.1 Roundup function: round up to nearest 0.1."""
        # Multiply by 10, ceiling, divide by 10
        int_val = int(value * 100_000)
        if int_val % 10_000 == 0:
            return int_val / 100_000
        return math.floor(int_val / 10_000 + 1) / 10.0

    @staticmethod
    def _severity(score: float) -> str:
        """Return the CVSS v3.1 qualitative severity rating."""
        if score == 0.0:
            return "None"
        elif score < 4.0:
            return "Low"
        elif score < 7.0:
            return "Medium"
        elif score < 9.0:
            return "High"
        else:
            return "Critical"

    @staticmethod
    def _parse_vector(vector_string: str) -> Dict[str, str]:
        """
        Parse a CVSS v3.x vector string into a dict of metric abbreviations.

        Handles both ``CVSS:3.1/AV:N/...`` and ``AV:N/...`` forms.
        """
        s = vector_string.strip()
        # Strip the CVSS:3.x/ prefix if present
        s = re.sub(r"^CVSS:[0-9.]+/", "", s, flags=re.IGNORECASE)

        metrics: Dict[str, str] = {}
        for part in s.split("/"):
            if ":" not in part:
                continue
            key, _, val = part.partition(":")
            metrics[key.upper()] = val.upper()

        # Validate required keys are present
        required = {"AV", "AC", "PR", "UI", "S", "C", "I", "A"}
        missing = required - metrics.keys()
        if missing:
            raise ValueError(
                f"CVSS vector is missing required metrics: {', '.join(sorted(missing))}. "
                f"Input was: '{vector_string}'"
            )
        return metrics


# ---------------------------------------------------------------------------
# DOCXExporter (Single Responsibility: Markdown-to-DOCX conversion only)
# ---------------------------------------------------------------------------


class DOCXExporter:
    """
    Convert a Markdown report string to a .docx file.

    Requires the ``python-docx`` package.  If it is not installed, all
    methods return ``None`` and log a warning rather than raising.

    Supported Markdown constructs:

    * ``# Title``      -> Document title paragraph (``Title`` style)
    * ``## Heading``   -> Heading 1
    * ``### Heading``  -> Heading 2
    * `` ``` ... ``` `` blocks -> ``Code`` style (or ``Normal`` + monospace)
    * ``---``          -> horizontal rule (paragraph border)
    * Regular text     -> Normal paragraph

    The document receives a "CONFIDENTIAL" header and a page-number footer.
    """

    def export(
        self,
        markdown_text: str,
        output_path: Union[str, Path],
    ) -> Optional[Path]:
        """
        Convert *markdown_text* to a DOCX file at *output_path*.

        Returns the written Path on success, or None if python-docx is
        unavailable or an error occurs.
        """
        try:
            from docx import Document                          # type: ignore
            from docx.shared import Pt, RGBColor              # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH     # type: ignore
            from docx.oxml.ns import qn                       # type: ignore
            from docx.oxml import OxmlElement                 # type: ignore
        except ImportError:
            log.warning(
                "python-docx is not installed; DOCX export skipped. "
                "Install with: pip install python-docx"
            )
            return None

        try:
            doc = Document()
            self._add_header_footer(doc)

            in_code_block = False
            code_lines: List[str] = []

            for line in markdown_text.splitlines():
                # Code block fence detection
                if line.strip().startswith("```"):
                    if in_code_block:
                        # Flush accumulated code lines
                        self._add_code_block(doc, "\n".join(code_lines))
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_lines.append(line)
                    continue

                # Heading levels
                if line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=2)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=1)
                elif line.startswith("# "):
                    p = doc.add_paragraph(line[2:].strip(), style="Title")
                elif line.strip() == "---":
                    self._add_horizontal_rule(doc)
                elif line.strip() == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)

            # Flush any unclosed code block
            if code_lines:
                self._add_code_block(doc, "\n".join(code_lines))

            out = Path(output_path)
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))
            log.info("DOCX report written to %s", out)
            return out

        except Exception as exc:
            log.error("DOCX export failed: %s", exc)
            return None

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _add_code_block(doc: Any, text: str) -> None:
        """Add a code-block paragraph using the Code style if available."""
        try:
            p = doc.add_paragraph(text, style="Code")
        except Exception:
            # Fallback: Normal style with Courier New
            from docx.shared import Pt  # type: ignore
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    @staticmethod
    def _add_horizontal_rule(doc: Any) -> None:
        """Add a paragraph that renders as a thin horizontal rule."""
        from docx.oxml.ns import qn   # type: ignore
        from docx.oxml import OxmlElement  # type: ignore

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _add_header_footer(doc: Any) -> None:
        """
        Insert a 'CONFIDENTIAL' header and page-number footer into every
        section of the document.
        """
        from docx.shared import Pt, RGBColor          # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn                    # type: ignore
        from docx.oxml import OxmlElement              # type: ignore

        section = doc.sections[0]

        # Header: "CONFIDENTIAL" centred in red
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = "CONFIDENTIAL"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.bold  = True
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Footer: "Page X" right-aligned using Word field codes
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.clear()

        run = fp.add_run("Page ")
        run.font.size = Pt(9)

        # Insert { PAGE } field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

_BASE_DIR     = Path(__file__).parent.parent
_SESSIONS_DIR = _BASE_DIR / "sessions"

# Module-level singletons for the new capabilities
_cvss_calc   = CVSSv3Calculator()
_docx_export = DOCXExporter()


class ReportGenerator:
    """
    Reads session artefacts and produces a Markdown pentest report.

    Parameters
    ----------
    sessions_dir : path to the sessions/ directory (default: <project>/sessions/)
    """

    def __init__(self, sessions_dir: str | Path = _SESSIONS_DIR) -> None:
        self.sdir = Path(sessions_dir)

    # ── Public API ────────────────────────────────────────────────────────────

    def generate(
        self,
        output_path: Optional[Union[str, Path]] = None,
        formats: Optional[List[str]] = None,
    ) -> Union[Path, List[Path]]:
        """
        Generate the penetration test report.

        Parameters
        ----------
        output_path : optional explicit output path.
            - If it ends in ``.docx``, the report is written as a DOCX file
              (requires python-docx) instead of Markdown.
            - If *None*, a timestamped ``.md`` file is created in
              ``sessions/``.
        formats : optional list of format strings, e.g. ``["md", "docx"]``.
            When provided, the report is written in all requested formats and
            a *list* of Paths is returned (one per format).  Supported values:
            ``"md"`` and ``"docx"``.  When *formats* is ``None`` (default),
            the behaviour is unchanged and a single Path is returned.

        Returns
        -------
        Path
            The path of the written file, when *formats* is None.
        List[Path]
            A list of written paths, when *formats* is a list (may include
            ``None`` entries if a format failed).
        """
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Build the Markdown content regardless of target format
        sections = [
            self._header(),
            self._section_scope(),
            self._section_findings(),
            self._section_credentials(),
            self._section_timeline(),
            self._section_objectives(),
            self._section_plan(),
            self._section_recommendations(),
            self._footer(),
        ]
        report_md = "\n\n".join(s for s in sections if s and s.strip())

        # --- Multi-format mode ---
        if formats is not None:
            results: List[Optional[Path]] = []
            base_stem = f"report_{ts}"
            for fmt in formats:
                fmt_lower = fmt.lower().strip()
                if fmt_lower == "md":
                    md_path = (
                        Path(output_path).with_suffix(".md")
                        if output_path
                        else self.sdir / f"{base_stem}.md"
                    )
                    md_path.parent.mkdir(parents=True, exist_ok=True)
                    md_path.write_text(report_md, encoding="utf-8")
                    log.info("Markdown report written to %s", md_path)
                    results.append(md_path)
                elif fmt_lower == "docx":
                    docx_path = (
                        Path(output_path).with_suffix(".docx")
                        if output_path
                        else self.sdir / f"{base_stem}.docx"
                    )
                    written = _docx_export.export(report_md, docx_path)
                    results.append(written)
                else:
                    log.warning("Unknown report format '%s'; skipping.", fmt)
            return results  # type: ignore[return-value]

        # --- Single-format mode (original behaviour, extended for .docx) ---
        if output_path is None:
            output_path = self.sdir / f"report_{ts}.md"
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        if out.suffix.lower() == ".docx":
            written = _docx_export.export(report_md, out)
            if written is None:
                # Fallback: write as Markdown with .md extension
                fallback = out.with_suffix(".md")
                fallback.write_text(report_md, encoding="utf-8")
                log.warning(
                    "DOCX export failed; Markdown fallback written to %s", fallback
                )
                return fallback
            return written

        out.write_text(report_md, encoding="utf-8")
        log.info("Report written to %s", out)
        return out

    # ── Sections ──────────────────────────────────────────────────────────────

    def _header(self) -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        facts = self._load_facts()
        hosts = list(facts.get("hosts", {}).keys()) if facts else []
        scope = ", ".join(hosts) if hosts else "unknown"
        return (
            "# Penetration Test Report\n\n"
            f"**Generated:** {now}  \n"
            f"**Framework:** LazyOwn RedTeam Suite  \n"
            f"**Scope:** {scope}  \n"
            f"**Classification:** CONFIDENTIAL\n\n"
            "---"
        )

    def _section_scope(self) -> str:
        facts = self._load_facts()
        if not facts:
            return "## 1. Scope & Discovery\n\n_No scan data available._"

        hosts = facts.get("hosts", {})
        lines = [
            "## 1. Scope & Discovery",
            "",
            f"**Hosts discovered:** {len(hosts)}",
            "",
            "| Host | OS | Open Services |",
            "|------|----|---------------|",
        ]
        for ip, info in hosts.items():
            services = info.get("services", {})
            if isinstance(services, dict):
                svc_list = [f"{port}/{info_s.get('name', '?')}" for port, info_s in services.items()]
            else:
                svc_list = [f"{s.get('port')}/{s.get('name','?')}" for s in (services or [])]
            os_hint = info.get("os_hint", info.get("os", "unknown"))
            lines.append(f"| {ip} | {os_hint} | {', '.join(svc_list) or 'none'} |")
        return "\n".join(lines)

    def _section_findings(self) -> str:
        facts = self._load_facts()
        if not facts:
            return ""

        vulns: List[dict] = []
        for ip, info in facts.get("hosts", {}).items():
            for v in info.get("vulnerabilities", []):
                vulns.append({"host": ip, **v})

        if not vulns:
            return "## 2. Findings\n\n_No vulnerabilities recorded._"

        lines = ["## 2. Findings", ""]
        for i, v in enumerate(vulns, 1):
            severity = v.get("severity", "INFO").upper()
            title    = v.get("title", v.get("name", f"Finding {i}"))
            desc     = v.get("description", v.get("detail", "No description."))
            evidence = v.get("evidence", v.get("output", ""))

            # CVSS v3.1 scoring if a vector or score is present
            cvss_line = ""
            cvss_raw = v.get("cvss", v.get("cvss_vector", ""))
            if cvss_raw:
                cvss_str = str(cvss_raw).strip()
                # Attempt to parse as a vector string; fall back if it looks
                # like a plain numeric score
                try:
                    if re.search(r"AV:", cvss_str, re.IGNORECASE):
                        score, label = _cvss_calc.calculate(cvss_str)
                        cvss_line = (
                            f"**CVSS v3.1:** {score} ({label})  \n"
                            f"**Vector:** `{cvss_str}`  "
                        )
                    else:
                        # Treat as a pre-computed numeric score
                        numeric = float(cvss_str)
                        label   = CVSSv3Calculator._severity(numeric)
                        cvss_line = f"**CVSS v3.1 Score:** {numeric:.1f} ({label})  "
                except Exception as exc:
                    log.debug("Could not parse CVSS field '%s': %s", cvss_str, exc)
                    cvss_line = f"**CVSS:** {cvss_str}  "

            lines += [
                f"### 2.{i} {title}",
                f"**Severity:** {severity}  ",
                f"**Host:** {v.get('host', 'N/A')}  ",
            ]
            if cvss_line:
                lines.append(cvss_line)
            lines.append(f"**Description:** {desc}")
            if evidence:
                lines.append(f"\n```\n{str(evidence)[:400]}\n```")
            lines.append("")
        return "\n".join(lines)

    def _section_credentials(self) -> str:
        cred_file = self.sdir / "credentials.txt"
        if not cred_file.exists():
            return ""
        creds = cred_file.read_text(encoding="utf-8", errors="replace").strip()
        if not creds:
            return ""
        count = len([l for l in creds.splitlines() if l.strip()])
        return (
            "## 3. Credentials Captured\n\n"
            f"**Total:** {count}\n\n"
            "```\n" + creds + "\n```"
        )

    def _section_timeline(self) -> str:
        events = self._load_jsonl(self.sdir / "events.jsonl")
        if not events:
            return ""

        lines = [
            "## 4. Timeline",
            "",
            "| Timestamp | Event Type | Detail |",
            "|-----------|------------|--------|",
        ]
        for ev in events[-60:]:
            ts     = str(ev.get("timestamp", ev.get("ts", "")))[:19]
            etype  = ev.get("type", ev.get("event_type", ""))
            detail = str(ev.get("data", ev.get("detail", ev.get("message", ""))))[:100]
            detail = detail.replace("|", "/")
            lines.append(f"| {ts} | {etype} | {detail} |")
        return "\n".join(lines)

    def _section_objectives(self) -> str:
        objs = self._load_jsonl(self.sdir / "objectives.jsonl")
        if not objs:
            return ""

        done    = [o for o in objs if o.get("status") in ("done", "completed")]
        pending = [o for o in objs if o.get("status") not in ("done", "completed")]

        lines = [
            "## 5. Objectives",
            "",
            f"**Completed:** {len(done)} / {len(objs)}",
            "",
        ]
        if done:
            lines.append("**Completed:**")
            for o in done:
                lines.append(f"- [x] {o.get('title', o.get('description', ''))}")
        if pending:
            lines.append("\n**Pending:**")
            for o in pending:
                lines.append(f"- [ ] {o.get('title', o.get('description', ''))}")
        return "\n".join(lines)

    def _section_plan(self) -> str:
        plan_file = self.sdir / "plan.txt"
        if not plan_file.exists():
            return ""
        plan = plan_file.read_text(encoding="utf-8", errors="replace").strip()
        if not plan:
            return ""
        return "## 6. Attack Plan\n\n```\n" + plan[:4000] + "\n```"

    def _section_recommendations(self) -> str:
        facts = self._load_facts()
        lines = ["## 7. Recommendations", ""]

        if not facts:
            lines.append("_Complete the engagement to generate recommendations._")
            return "\n".join(lines)

        SERVICE_RECS: Dict[str, str] = {
            "http":  "Perform web application assessment: directory brute-force, auth bypass, injection, misconfigurations.",
            "https": "Perform web application assessment over TLS: cert validity, HSTS, injection, auth bypass.",
            "smb":   "Test for null session, EternalBlue (MS17-010), SMB relay, credential brute-force.",
            "ssh":   "Check for outdated version vulnerabilities, user enumeration, brute-force, weak keys.",
            "ftp":   "Test anonymous login, brute-force, cleartext credential sniffing.",
            "ldap":  "Enumerate domain objects, AS-REP roast, Kerberoast, ACL abuse.",
            "ldaps": "Enumerate domain objects over LDAPS, AS-REP roast, Kerberoast, ACL abuse.",
            "mssql": "Test sa account, xp_cmdshell, credential stuffing, UNC path injection.",
            "mysql": "Test for blank root password, file read/write via LOAD DATA INFILE.",
            "rdp":   "Test BlueKeep (CVE-2019-0708), NLA bypass, credential brute-force.",
            "winrm": "Test credential brute-force; if credentials known, attempt lateral movement.",
        }

        added = False
        for ip, info in facts.get("hosts", {}).items():
            services = info.get("services", {})
            if isinstance(services, dict):
                svc_items = [(port, s.get("name", "")) for port, s in services.items()]
            else:
                svc_items = [(s.get("port", "?"), s.get("name", "")) for s in (services or [])]

            for port, name in svc_items:
                name_l = name.lower().split("/")[0].strip()
                rec = SERVICE_RECS.get(name_l)
                if rec:
                    lines.append(f"- **{ip}:{port}** (`{name_l}`): {rec}")
                    added = True

        if not added:
            lines.append("_No service-specific recommendations generated from current fact data._")
        return "\n".join(lines)

    def _footer(self) -> str:
        return (
            "---\n\n"
            "_This report was auto-generated by LazyOwn. "
            "Review and validate all findings before submission._"
        )

    # ── Data loaders ──────────────────────────────────────────────────────────

    def _load_facts(self) -> Dict[str, Any]:
        for fname in ("policy_facts.json", "facts.json"):
            p = self.sdir / fname
            if p.exists():
                try:
                    raw = json.loads(p.read_text(encoding="utf-8"))
                    # Normalise: some fact stores use top-level hosts dict, others wrap it
                    if "hosts" in raw and isinstance(raw["hosts"], dict):
                        return raw
                    # Treat the whole dict as hosts only if all values are dicts (host records)
                    host_like = {k: v for k, v in raw.items() if isinstance(v, dict)}
                    if host_like:
                        return {"hosts": host_like}
                    return {}
                except Exception as exc:
                    log.warning("Could not load %s: %s", p, exc)
        return {}

    def _load_jsonl(self, path: Path) -> List[Dict]:
        if not path.exists():
            return []
        results = []
        for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                results.append(json.loads(line))
            except Exception:
                pass
        return results


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")

    p = argparse.ArgumentParser(description="LazyOwn Report Generator")
    p.add_argument("--sessions", default=str(_SESSIONS_DIR),
                   help="Path to sessions directory")
    p.add_argument("--output",   default=None,
                   help="Output file path (default: sessions/report_<ts>.md)")
    p.add_argument("--quiet",    action="store_true")
    args = p.parse_args()

    if args.quiet:
        logging.disable(logging.WARNING)

    rg  = ReportGenerator(sessions_dir=args.sessions)
    out = rg.generate(output_path=args.output)
    print(out)
